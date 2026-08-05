"""
Corporate Proposal Document Export Engine for TEKLİF-Sim (v3.0.0).
Generates official EASA Part 21 format downloadable PDF (via ReportLab)
and Word DOCX (via python-docx) proposal documents.
"""

import io
from typing import Dict, Any, Optional

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.logger import logger


def generate_pdf_proposal(proposal_data: Dict[str, Any]) -> bytes:
    """
    Generates a corporate PDF proposal document using ReportLab.
    Returns PDF content as bytes.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    
    styles = getSampleStyleSheet()
    
    # Custom Corporate Styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#0F172A'),
        spaceAfter=6
    )
    subtitle_style = ParagraphStyle(
        'DocSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#64748B'),
        spaceAfter=12
    )
    section_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=16,
        textColor=colors.HexColor('#C8102E'),
        spaceBefore=10,
        spaceAfter=6
    )
    body_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=14,
        textColor=colors.HexColor('#1E293B')
    )

    story = []

    # Title & Corporate Header
    story.append(Paragraph("AeroDesign Engineering | Proposal Quotation", title_style))
    story.append(Paragraph("EASA Part 21J Design Organisation Approval (DOA) Ref: EASA.21J.999", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#C8102E'), spaceAfter=12))

    # Executive Overview Table
    customer = proposal_data.get("customer_name") or "Flagship Air"
    aircraft = proposal_data.get("aircraft_type") or "A320-200"
    fleet = proposal_data.get("fleet_size") or 1
    mod_type = proposal_data.get("modification_type") or "cabin"
    currency = proposal_data.get("currency") or "USD"
    total_price = proposal_data.get("total_price_formatted") or f"${proposal_data.get('total_price_usd', 0):,.2f}"

    overview_data = [
        [Paragraph("<b>Customer Airline:</b>", body_style), Paragraph(str(customer), body_style),
         Paragraph("<b>Aircraft Model:</b>", body_style), Paragraph(str(aircraft), body_style)],
        [Paragraph("<b>Fleet Size:</b>", body_style), Paragraph(f"{fleet} Aircraft", body_style),
         Paragraph("<b>Mod Category:</b>", body_style), Paragraph(str(mod_type).upper(), body_style)],
        [Paragraph("<b>Proposal Currency:</b>", body_style), Paragraph(str(currency), body_style),
         Paragraph("<b>Total Quote:</b>", body_style), Paragraph(f"<b>{total_price}</b>", body_style)],
    ]
    t_overview = Table(overview_data, colWidths=[1.3*inch, 2.2*inch, 1.3*inch, 2.2*inch])
    t_overview.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#F8FAFC')),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#CBD5E1')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t_overview)
    story.append(Spacer(1, 14))

    # Project Scope & Regulatory Basis
    story.append(Paragraph("1. Certification & Scope Summary", section_style))
    scope_text = proposal_data.get("scope_text") or "Engineering design modification, STC compliance verification, and Part 21 package preparation."
    story.append(Paragraph(f"<b>Scope of Work:</b> {scope_text}", body_style))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Certification Basis:</b> CS-25 (Large Aeroplanes) / CS-23 (Normal Aeroplanes). Complies with EASA Part 21.A.239.", body_style))
    story.append(Spacer(1, 12))

    # Itemized Financial Cost Breakdown Table
    story.append(Paragraph("2. Financial & Engineering Cost Kırılımı", section_style))
    
    quote_data = proposal_data.get("quote_breakdown", {})
    labor_val = quote_data.get("base_labor_cost_adjusted", quote_data.get("base_labor_cost", 0.0))
    contingency_val = quote_data.get("contingency", 0.0)
    materials_val = quote_data.get("material_allowance", 0.0)
    testing_val = quote_data.get("testing_fee", 0.0)
    volume_disc_val = quote_data.get("volume_discount", 0.0)

    breakdown_table_data = [
        [Paragraph("<b>Cost Line Item</b>", body_style), Paragraph("<b>Description</b>", body_style), Paragraph("<b>Amount</b>", body_style)],
        [Paragraph("Engineering Labor", body_style), Paragraph("DOA design, stress, avionics & cert hours", body_style), Paragraph(f"${labor_val:,.2f}", body_style)],
        [Paragraph("Risk Contingency", body_style), Paragraph("AACE risk-adjusted engineering contingency", body_style), Paragraph(f"${contingency_val:,.2f}", body_style)],
        [Paragraph("Material Allowance", body_style), Paragraph("Kitting & structural hardware per fleet", body_style), Paragraph(f"${materials_val:,.2f}", body_style)],
        [Paragraph("Testing & Qualification", body_style), Paragraph("CS-25.1309 / DO-160G testing fees", body_style), Paragraph(f"${testing_val:,.2f}", body_style)],
    ]
    if volume_disc_val > 0:
        breakdown_table_data.append([
            Paragraph("Volume Discount", body_style),
            Paragraph(f"Fleet volume discount ({quote_data.get('volume_discount_rate', 0)*100:.0f}%)", body_style),
            Paragraph(f"-${volume_disc_val:,.2f}", body_style)
        ])

    t_breakdown = Table(breakdown_table_data, colWidths=[2.2*inch, 3.5*inch, 1.3*inch])
    t_breakdown.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0F172A')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#CBD5E1')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
    ]))
    story.append(t_breakdown)
    story.append(Spacer(1, 14))

    # Risk Confidence Intervals (Monte-Carlo)
    if "p10_p50_p90" in proposal_data:
        mc = proposal_data["p10_p50_p90"]
        story.append(Paragraph("3. Monte-Carlo Risk & Confidence Intervals", section_style))
        mc_data = [
            [Paragraph("<b>Confidence Level</b>", body_style), Paragraph("<b>Estimated Hours</b>", body_style), Paragraph("<b>Estimated Cost</b>", body_style)],
            [Paragraph("P10 (Optimistic)", body_style), Paragraph(f"{mc['p10']['manhours']:.0f} hrs", body_style), Paragraph(f"${mc['p10']['cost']:,.2f}", body_style)],
            [Paragraph("P50 (Expected Median)", body_style), Paragraph(f"{mc['p50']['manhours']:.0f} hrs", body_style), Paragraph(f"${mc['p50']['cost']:,.2f}", body_style)],
            [Paragraph("P90 (Conservative Worst-Case)", body_style), Paragraph(f"{mc['p90']['manhours']:.0f} hrs", body_style), Paragraph(f"${mc['p90']['cost']:,.2f}", body_style)],
        ]
        t_mc = Table(mc_data, colWidths=[2.5*inch, 2.0*inch, 2.5*inch])
        t_mc.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#334155')),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#CBD5E1')),
            ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
            ('TOPPADDING', (0,0), (-1,-1), 5),
            ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ]))
        story.append(t_mc)
        story.append(Spacer(1, 14))

    # Terms & Authorization
    story.append(Paragraph("4. Terms & Sign-off", section_style))
    story.append(Paragraph("Proposal valid for 60 calendar days from issue date. Price excludes regional VAT/taxes unless specified.", body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_docx_proposal(proposal_data: Dict[str, Any]) -> bytes:
    """
    Generates a corporate Word DOCX proposal document using python-docx.
    Returns DOCX content as bytes.
    """
    doc = docx.Document()
    
    # Title
    h1 = doc.add_heading("AeroDesign Engineering | Official Proposal", level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p_sub = doc.add_paragraph("EASA Part 21J Design Organisation Approval (DOA) Ref: EASA.21J.999")
    p_sub.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    p_sub.runs[0].font.size = Pt(10)
    
    doc.add_paragraph() # Spacer

    # Executive Overview
    customer = proposal_data.get("customer_name") or "Flagship Air"
    aircraft = proposal_data.get("aircraft_type") or "A320-200"
    fleet = proposal_data.get("fleet_size") or 1
    total_price = proposal_data.get("total_price_formatted") or f"${proposal_data.get('total_price_usd', 0):,.2f}"

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).text = f"Customer Airline: {customer}"
    table.cell(0, 1).text = f"Aircraft Model: {aircraft}"
    table.cell(1, 0).text = f"Fleet Size: {fleet} Aircraft"
    table.cell(1, 1).text = f"Mod Type: {proposal_data.get('modification_type', 'cabin').upper()}"
    table.cell(2, 0).text = f"Currency: {proposal_data.get('currency', 'USD')}"
    table.cell(2, 1).text = f"Total Quote: {total_price}"

    doc.add_heading("1. Certification & Scope Summary", level=2)
    doc.add_paragraph(f"Scope of Work: {proposal_data.get('scope_text', 'Engineering design modification and STC approval package.')}")
    doc.add_paragraph("Certification Basis: CS-25 (Large Aeroplanes) / CS-23 (Normal Aeroplanes). Complies with EASA Part 21.A.239.")

    doc.add_heading("2. Financial Kırılım & Terms", level=2)
    doc.add_paragraph("Quotations are generated deterministically according to AeroDesign rate cards and EASA DOA guidelines.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
