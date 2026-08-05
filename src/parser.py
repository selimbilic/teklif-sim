"""
Document Parsing Module for TEKLİF-Sim (v3.0.0).
Extracts raw text and tables from uploaded PDF, Excel, Word, and Text files.
"""

import io
import os
from typing import Optional, Union
from src.logger import logger

def parse_pdf(file_bytes: bytes) -> str:
    """Extracts text from a PDF file using pypdf or pdfplumber."""
    extracted_text = []
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                extracted_text.append(f"--- Page {page_num + 1} ---\n{text}")
    except Exception as e:
        logger.warning(f"pypdf extraction failed or partial: {e}")
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        extracted_text.append(f"--- Page {page_num + 1} ---\n{text}")
        except Exception as e2:
            logger.error(f"pdfplumber extraction failed: {e2}")

    return "\n\n".join(extracted_text)


def parse_excel(file_bytes: bytes) -> str:
    """Extracts tables and text from Excel spreadsheets (.xlsx, .xls)."""
    try:
        import pandas as pd
        excel_file = pd.ExcelFile(io.BytesIO(file_bytes))
        sheet_texts = []
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            if not df.empty:
                try:
                    table_str = df.to_markdown(index=False)
                except Exception:
                    table_str = df.to_csv(index=False)
                sheet_texts.append(f"### Sheet: {sheet_name}\n" + table_str)
        return "\n\n".join(sheet_texts)
    except Exception as e:
        logger.error(f"Excel parsing failed: {e}")
        return f"[Excel Parsing Error: {e}]"


def parse_word(file_bytes: bytes) -> str:
    """Extracts text and tables from Microsoft Word documents (.docx)."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_vals = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_vals:
                    full_text.append(" | ".join(row_vals))
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Word parsing failed: {e}")
        return f"[Word Parsing Error: {e}]"


def parse_uploaded_file(file_name: str, file_bytes: bytes) -> str:
    """
    Main entry point for parsing an uploaded document by extension.
    Returns cleaned markdown/text content.
    """
    ext = os.path.splitext(file_name)[1].lower()
    
    if ext == ".pdf":
        return parse_pdf(file_bytes)
    elif ext in [".xlsx", ".xls"]:
        return parse_excel(file_bytes)
    elif ext in [".docx", ".doc"]:
        return parse_word(file_bytes)
    elif ext in [".txt", ".md", ".csv", ".json", ".log"]:
        try:
            return file_bytes.decode("utf-8", errors="replace")
        except Exception as e:
            return f"[Text Decoding Error: {e}]"
    else:
        logger.warning(f"Unsupported file format for text extraction: {ext}")
        return f"[Unsupported file type: {ext}]"
